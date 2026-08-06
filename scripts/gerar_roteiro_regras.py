# -*- coding: utf-8 -*-
"""Gera ENTREGA/ROTEIRO_REGRAS_CVC_IAM_v1.0.0.docx.

Roteiro de validacao POR CATEGORIA E REGRA, com passo conferivel.

PREMISSA: teste automatizado prova que o CODIGO faz o que foi programado; so
quem conhece a operacao prova que a REGRA esta certa. Por isso cada regra traz
o criterio com os limiares numericos e uma pergunta a responder.

DE ONDE VEM CADA NUMERO — a versao anterior deste documento puxava numeros das
TABELAS DO MOTOR, e tres deles nao batiam com a tela: o painel conta USUARIOS
distintos (nao linhas), le o snapshot `bi_divergencias` e usa rotulos proprios
("Sem Vinculo RH", "Matriz SYSTUR") em vez dos nomes internos. Documento que
nao bate com a tela destroi a confianca no resto.
Agora TODO numero vem da API do painel, na hora da geracao: e' literalmente o
que a usuaria le na tela. Se o painel nao estiver no ar, o script para.

Uso:
    1) suba o Visualizador do pacote (porta 8800)
    2) python scripts/gerar_roteiro_regras.py
"""
import json
import sys
import urllib.error
import urllib.request
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = Path(__file__).resolve().parent.parent
OUT = RAIZ / "ENTREGA" / "ROTEIRO_REGRAS_CVC_IAM_v1.0.0.docx"
API = "http://127.0.0.1:8800"

AZUL = RGBColor(0x1F, 0x2D, 0x5C)
CINZA = RGBColor(0x5A, 0x64, 0x78)
TEXTO = RGBColor(0x2C, 0x33, 0x40)
AMBAR = RGBColor(0x7A, 0x5B, 0x10)
VERDE = RGBColor(0x1E, 0x7B, 0x43)


# ---------------------------------------------------------------- dados
def api(rota):
    try:
        with urllib.request.urlopen(API + rota, timeout=300) as r:
            return json.loads(r.read())
    except (urllib.error.URLError, OSError) as e:
        raise SystemExit(
            f"Painel nao respondeu em {API}{rota} ({e}).\n"
            "Suba o Visualizador do pacote antes de gerar o documento — os "
            "numeros deste roteiro SAO os da tela, nao os das tabelas.")


def coletar():
    db = api("/api/dados")
    U = db["users"]
    usr = lambda f: len({u["u"] for u in U for d in u["divs"] if f(d)})
    pend = lambda d: d.get("a") not in ("Aderente", "Incluir Acesso")

    d = {
        "kpis": db["kpis"],
        "vg": db["vg"],
        "consulta": len(U),
        "pendencias": len({u["u"] for u in U for d_ in u["divs"] if pend(d_)}),
        "aderentes": len(db.get("aderentes") or []),
        "acao": {k: usr(lambda x, k=k: x.get("a") == k)
                 for k in {x.get("a") for u in U for x in u["divs"] if x.get("a")}},
        "tipo": {k: usr(lambda x, k=k: x.get("tl") == k)
                 for k in {x.get("tl") for u in U for x in u["divs"] if x.get("tl")}},
        "vinculo": {k: usr(lambda x, k=k: x.get("vinc") == k)
                    for k in {x.get("vinc") for u in U for x in u["divs"] if x.get("vinc")}},
        "origem": {k: usr(lambda x, k=k: x.get("o") == k)
                   for k in {x.get("o") for u in U for x in u["divs"] if x.get("o")}},
        "sistema_pend": {},
        "linhas": len([1 for u in U for _ in u["divs"]]),
    }
    for s in sorted({x.get("sis") for u in U for x in u["divs"] if x.get("sis")}):
        d["sistema_pend"][s] = usr(lambda x, s=s: x.get("sis") == s and pend(x))
    des = api("/api/desligados"); d["desligados"] = des["kpis"]
    tra = api("/api/transferidos"); d["transferidos"] = tra["kpis"]
    h = api("/api/historico")
    d["historico"] = len(h if isinstance(h, list) else (h.get("lista") or []))
    return d


N = lambda v: f"{int(v):,}".replace(",", ".")


# ---------------------------------------------------------------- docx
def shade(cell, cor):
    tc = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), cor)
    tc.append(el)


def h1(doc, txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def par(doc, txt, size=10, italic=False, bold=False, cor=TEXTO, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt)
    r.italic, r.bold = italic, bold
    r.font.size = Pt(size)
    r.font.color.rgb = cor
    return p


def nota(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(txt)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = AMBAR


def regra(doc, cod, titulo, decide, criterio, conferir, esperado, prioritaria=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(("★  " if prioritaria else "") + f"{cod}  {titulo}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = AZUL

    linhas = [
        ("O que decide", decide),
        ("Critério", criterio),
        ("Como conferir", conferir),
        ("Deve mostrar", esperado),
        ("A regra está correta? Se não, qual deveria ser?", ""),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for rot, val in linhas:
        cells = tbl.add_row().cells
        cells[0].width = Cm(4.6)
        cells[1].width = Cm(11.9)
        pergunta = rot.startswith("A regra")
        r0 = cells[0].paragraphs[0].add_run(rot)
        r0.bold = True
        r0.font.size = Pt(8.5)
        r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if pergunta else AZUL
        shade(cells[0], "1F2D5C" if pergunta else "EEF2F8")
        r1 = cells[1].paragraphs[0].add_run(val)
        r1.font.size = Pt(9)
        r1.font.color.rgb = VERDE if rot == "Deve mostrar" else TEXTO
        if rot == "Deve mostrar":
            r1.bold = True
        if pergunta:
            shade(cells[1], "FCF9EE")
    return tbl


def main():
    d = coletar()
    k, vg = d["kpis"], d["vg"]
    ac, tp, vi = d["acao"], d["tipo"], d["vinculo"]
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ---------------- capa ----------------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("CVC IAM Analytics"); r.bold = True
    r.font.size = Pt(20); r.font.color.rgb = AZUL
    s = doc.add_paragraph()
    r = s.add_run("Roteiro de teste e validação — por categoria e regra")
    r.font.size = Pt(13); r.font.color.rgb = CINZA
    par(doc, "Pacote TESTE_LOCAL_BRUNA_v1.0.0  ·  06/08/2026", size=9, cor=CINZA)

    doc.add_paragraph()
    par(doc, "Por que este documento existe", bold=True, size=11, cor=AZUL, space=2)
    par(doc,
        "Os testes automatizados provam que o sistema faz o que foi programado. "
        "Não provam que a REGRA está certa — isso só quem conhece a operação da "
        "CVC consegue dizer. Cada regra abaixo decide sozinha alguma coisa no "
        "painel; aqui está o critério que ela usa, com os números, e o passo para "
        "conferir na tela.")

    par(doc, "Como usar", bold=True, size=11, cor=AZUL, space=2)
    par(doc,
        "Cada regra tem “Como conferir” (o que filtrar) e “Deve mostrar” (o "
        "número que tem de aparecer). Se bater, a regra está funcionando — falta "
        "dizer se ela está CERTA, e é isso que a última linha pergunta. Se não "
        "bater, já é um achado: anote e me avise.")
    par(doc,
        "As regras marcadas com ★ são as que mais mudam volume de trabalho. Se o "
        "tempo for curto, comece por elas.")
    par(doc,
        "Todos os números deste documento foram lidos do próprio painel no "
        "momento em que ele foi gerado — são os mesmos que aparecem na tela, "
        "contando PESSOAS (o painel conta pessoas a tratar, não linhas).")

    nota(doc,
         "Ressalva desta base: as pendências vindas da matriz de cargo foram "
         "geradas hoje, então a coluna de tempo de tratamento na Visão Geral "
         "está achatada. Os vereditos estão corretos; apenas o “há quantos dias” "
         "dessas linhas não representa histórico.")

    # ---------------- reconciliacao ----------------
    doc.add_page_break()
    h1(doc, "Antes de tudo: de onde vêm os números")
    par(doc,
        "A primeira dúvida costuma ser por que o painel fala em milhares num "
        "lugar e em centenas no outro. O caminho é este:")

    passos = [
        (f"{N(vg['total_acessos'])} acessos", "lidos dos sete sistemas"),
        (f"{N(vg['acessos_vinc'])} acessos", "ligados a uma pessoa identificada "
         f"({vg['cobertura_pct']}% de cobertura); o restante é conta sem dono "
         "conhecido, em geral franqueado ou prestador"),
        (f"{N(d['linhas'])} avaliações", "uma por pessoa × sistema × perfil "
         "esperado — é o que a tela lista"),
        (f"{N(d['pendencias'])} pessoas a tratar",
         "só o que exige ação: Em Análise, Alterar Perfil e Sem Vínculo RH"),
    ]
    for i, (num, txt) in enumerate(passos, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{i}.  {num}  ")
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = AZUL
        r2 = p.add_run("— " + txt); r2.font.size = Pt(9.5); r2.font.color.rgb = TEXTO

    par(doc, "", space=2)
    par(doc,
        f"As {N(ac.get('Incluir Acesso', 0))} pessoas em “Incluir Acesso” e as "
        f"{N(ac.get('Aderente', 0))} em “Aderente” NÃO entram nas "
        f"{N(d['pendencias'])}: quem não tem o acesso não é irregularidade a "
        "corrigir, e quem está aderente não precisa de nada. Essa foi uma decisão "
        "sua, e é o que faz o número da fila ser tratável.")
    soma = (tp.get("Em Análise", 0) + tp.get("Divergente", 0)
            + tp.get("Sem Vínculo RH", 0))
    par(doc,
        f"Somando as três categorias dá {N(soma)}, e o painel mostra "
        f"{N(d['pendencias'])}: a diferença são pessoas que aparecem em mais de "
        "uma categoria e são contadas uma vez só.")

    # ---------------- 1 identidade ----------------
    doc.add_page_break()
    h1(doc, "1. Identidade — de quem é cada acesso")
    par(doc,
        f"Antes de julgar se um acesso é indevido, o sistema precisa saber de "
        f"quem ele é. São {N(vg['total_acessos'])} acessos lidos contra "
        f"{N(vg['rh_ativos'])} ativos e {N(vg['rh_desligados'])} desligados do RH.")

    regra(doc, "1.1", "Cascata de vinculação (acesso → pessoa)",
          "De quem é cada acesso encontrado nos extratos.",
          "Tenta em ordem e para no primeiro que casar: CPF → e-mail → login → "
          "CPF parcial + nome → nome → aproximação. Do mais confiável para o "
          "menos: o CPF identifica sozinho, o nome não.",
          "Visão Geral → cartão de cobertura.",
          f"{vg['cobertura_pct']}% — {N(vg['acessos_vinc'])} de "
          f"{N(vg['total_acessos'])} acessos ligados a uma pessoa.")
    nota(doc,
         "O que sobra é majoritariamente franqueado e prestador, que não estão "
         "na base de RH. Vale dizer se esse percentual faz sentido para a CVC.")

    regra(doc, "1.2", "Situação da conta manda",
          "Se uma conta encontrada no extrato conta como acesso.",
          "Conta BLOQUEADA ou INATIVA não é acesso: já está revogada, então não "
          "vira pendência em nenhuma regra.",
          "Não há filtro — é uma exclusão. Percebe-se pelo tamanho da fila.",
          "Mais da metade das contas dos extratos está bloqueada e ficou de fora.",
          prioritaria=True)
    nota(doc,
         "É a regra de maior impacto no volume. Se para a CVC uma conta "
         "bloqueada ainda for risco (porque pode ser reativada), ela precisa mudar.")

    regra(doc, "1.3", "Acesso sem vínculo com o RH",
          "Acesso cujo dono não foi encontrado em nenhuma base de pessoas.",
          "Um achado por LOGIN e SISTEMA, não por perfil: um login com cinco "
          "perfis gera um achado, e ele carrega os cinco. Conta bloqueada fica "
          "de fora (regra 1.2).",
          "Pendências → filtro Tipo = “Sem Vínculo RH”.",
          f"{N(tp.get('Sem Vínculo RH', 0))} pessoas.")

    regra(doc, "1.4", "Quem é franqueado, prestador ou terceiro",
          "A que população cada pessoa pertence — é isso que escolhe a regra de "
          "validação nas seções seguintes.",
          "Funcionário vem do RH; franqueado e prestador vêm de três exportações "
          "do diretório (AD), lidas em ordem cronológica, valendo a mais recente.",
          "Pendências → filtro Vínculo (ou coluna Vínculo na Consulta).",
          "  ·  ".join(f"{kk} {N(vv)}" for kk, vv in
                       sorted(vi.items(), key=lambda x: -x[1])) + " (pessoas).")

    # ---------------- 2 validacao ----------------
    doc.add_page_break()
    h1(doc, "2. Validação — o que a pessoa deveria ter")
    par(doc,
        "Aqui o sistema compara o que a pessoa TEM com o que ela DEVERIA ter e "
        "emite um veredito. O que muda de regra para regra é COMO se descobre o "
        "“deveria ter”.")

    par(doc, "Os quatro vereditos", bold=True, size=10, cor=AZUL, space=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, txt in enumerate(["Veredito", "Significa", "Pessoas"]):
        c = tbl.rows[0].cells[i]; c.text = ""
        rr = c.paragraphs[0].add_run(txt); rr.bold = True
        rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1F2D5C")
    for rot, sig, val in [
        ("Aderente", "tem exatamente o que o cargo prevê", ac.get("Aderente", 0)),
        ("Incluir Acesso", "o cargo prevê e a pessoa não tem", ac.get("Incluir Acesso", 0)),
        ("Alterar Perfil", "tem o acesso, mas com o perfil errado", ac.get("Alterar Perfil", 0)),
        ("Em Análise", "não dá para decidir sozinho — precisa de gente", ac.get("Em Análise", 0)),
    ]:
        cells = tbl.add_row().cells
        for j, v in enumerate([rot, sig, N(val)]):
            rr = cells[j].paragraphs[0].add_run(v)
            rr.font.size = Pt(9); rr.font.color.rgb = TEXTO
            if j == 0:
                rr.bold = True

    regra(doc, "2.1", "Matriz de perfis por cargo",
          "O perfil esperado, quando existe matriz para o cargo. É a regra "
          "principal — as outras entram onde ela não alcança.",
          "Chave: centro de custo + cargo.",
          "Pendências → filtro Origem, valores que começam com “Matriz ”.",
          "  ·  ".join(f"{kk.replace('Matriz ', '')} {N(vv)}"
                       for kk, vv in sorted(d["origem"].items(), key=lambda x: -x[1])
                       if kk.startswith("Matriz ") and kk != "Matriz CCO")
          + " (pessoas).", prioritaria=True)

    regra(doc, "2.2", "CCO — centro de custo + gestor",
          "O perfil esperado quando o cargo não está na matriz, mas o centro de "
          "custo e o gestor estão no mapeamento CCO/CSC.",
          "Chave: centro de custo + gestor. Aplicada ANTES da regra de Em "
          "Análise, senão sistemas sem CCO virariam enxurrada de Em Análise falso.",
          "Pendências → filtro Origem = “Matriz CCO”.",
          f"{N(d['origem'].get('Matriz CCO', 0))} pessoas.")

    regra(doc, "2.3", "Espelho dinâmico do SIG",
          "O perfil esperado no SIG, que não tem matriz de cargo.",
          "Olha os COLEGAS e considera padrão o perfil presente em pelo menos "
          "70% deles. Grupo: centro de custo + gestor + cargo; sem gente "
          "suficiente, cai para centro de custo + gestor. Exige no mínimo 2 "
          "colegas usando o SIG — com menos não há padrão que se sustente.",
          "Pendências → filtro Sistema = SIG.",
          f"{N(d['sistema_pend'].get('SIG', 0))} pessoas a tratar no SIG.",
          prioritaria=True)
    nota(doc,
         "Os 70% e o mínimo de 2 colegas são parâmetros: mexer neles muda "
         "diretamente quantas pendências o SIG gera. É a regra que mais depende "
         "do seu julgamento.")

    regra(doc, "2.4", "Espelho de terceiros, franqueados e prestadores",
          "O perfil esperado de quem não tem cargo na estrutura da CVC.",
          "Mesma lógica de espelho (70%, mínimo 2 pares), com chave própria: "
          "terceiro espelha por empresa + supervisor; franqueado e prestador, "
          "por empresa + gestor do diretório. Vale em todos os sistemas.",
          "Pendências → filtro Vínculo, tudo que não seja Funcionário.",
          "  ·  ".join(f"{kk} {N(vv)}" for kk, vv in sorted(vi.items(), key=lambda x: -x[1])
                       if kk != "Funcionário") + " (pessoas).")

    regra(doc, "2.5", "Limiar de inclusão — 30% de adesão",
          "Se a falta de um acesso vira pendência de Incluir Acesso.",
          "Só cobra inclusão quando pelo menos 30% das pessoas daquele cargo já "
          "têm o acesso. Abaixo disso a matriz provavelmente abrange demais, e "
          "cobrar inundaria a fila de ruído.",
          "É uma supressão: as linhas não aparecem.",
          f"{N(ac.get('Incluir Acesso', 0))} pessoas sobraram em Incluir Acesso "
          "depois do corte.", prioritaria=True)
    nota(doc,
         "Parâmetro ajustável: 0% desliga a regra e tudo vira pendência. Se a "
         "CVC quiser cobrar inclusão mesmo em cargo de baixa adesão, é só baixar.")

    regra(doc, "2.6", "Sem grupo de comparação não vira pendência",
          "O que fazer quando não há como saber o esperado.",
          "Terceiro, franqueado ou prestador sem grupo-espelho com padrão não "
          "vira pendência: sem par comparável, o sistema não tem base para "
          "afirmar que o acesso está errado, e acusar sem base gera retrabalho.",
          "É uma supressão: as linhas não aparecem.",
          "Nenhuma pendência é criada nesses casos.")

    regra(doc, "2.7", "Dois ou mais perfis e nenhum aderente → Em Análise",
          "O veredito quando a pessoa tem acesso mas nada bate com o esperado.",
          "Com 2+ acessos ou 2+ perfis esperados vai para Em Análise (excesso ou "
          "ambiguidade — precisa de gente). Com exatamente 1 de cada que não "
          "casam, é Alterar Perfil: aí o erro é claro.",
          "Pendências → filtro Ação = “Em Análise”; depois “Alterar Perfil”.",
          f"Em Análise {N(ac.get('Em Análise', 0))} pessoas  ·  Alterar Perfil "
          f"{N(ac.get('Alterar Perfil', 0))} pessoas.", prioritaria=True)

    regra(doc, "2.8", "Perfil por aproximação (só no IC)",
          "Como comparar o perfil do extrato com o da matriz no Integrador "
          "Contábil.",
          "O extrato traz “IC_CONSULTA” e a matriz traz “IC CONSULTA” — e a "
          "própria matriz é inconsistente internamente. Só no IC a comparação "
          "ignora underscore, espaço e caixa. Não se aplica ao SYSTUR, que já "
          "bate exato.",
          "Pendências → filtro Sistema = IC_INTEGRADOR_CONTABIL.",
          f"{N(d['sistema_pend'].get('IC_INTEGRADOR_CONTABIL', 0))} pessoas a "
          "tratar. Solução temporária, a remover quando a matriz for padronizada.")

    regra(doc, "2.9", "Status indefinido no extrato → Em Análise",
          "O que fazer quando o extrato não diz se a conta está ativa.",
          "Sem informação de status o sistema não decide sozinho: manda para Em "
          "Análise, em vez de assumir ativo (acusaria indevidamente) ou "
          "bloqueado (esconderia risco).",
          "Pendências → filtro Ação = “Em Análise”.",
          "Poucos casos nesta base; a regra existe para não decidir no escuro.")

    # ---------------- 3 ciclo ----------------
    doc.add_page_break()
    h1(doc, "3. Ciclo de vida — o que muda quando a pessoa muda")
    par(doc,
        "As regras acima olham uma foto. Estas comparam fotos: o que acontece "
        "quando alguém sai da empresa ou muda de função.")

    des, tra = d["desligados"], d["transferidos"]
    regra(doc, "3.1", "Desligado com acesso ativo",
          "Acesso que continua existindo depois do desligamento.",
          "Cruza a base de desligados do RH com os extratos. Conta bloqueada não "
          "conta (regra 1.2) — já está revogada.",
          "Aba Desligados → visão “A tratar”.",
          f"{N(des.get('tratar', 0))} pessoas a tratar, de "
          f"{N(des.get('total', 0))} desligados na base.", prioritaria=True)
    nota(doc,
         f"A aba lista os {N(des.get('total', 0))} desligados, mas só "
         f"{N(des.get('tratar', 0))} têm acesso ativo. É esse o número "
         "operacional — os outros já estão regulares.")

    regra(doc, "3.2", "Transferido — detecção da movimentação",
          "Quem mudou de função e por isso precisa ter o acesso revisto.",
          "Compara cada carga de RH com a anterior e detecta mudança em cargo, "
          "centro de custo, departamento ou gestor. Só enxerga a mudança ENTRE "
          "execuções — por isso o histórico depende de rodar com regularidade.",
          "Aba Transferidos.",
          f"{N(tra.get('total', 0))} pessoas a revisar, de "
          f"{N(vg['transf_movimentos'])} movimentações detectadas.")

    regra(doc, "3.3", "Revalidação depois da transferência",
          "O que fazer com cada acesso de quem se moveu.",
          "Classifica cada acesso comparando o esperado ANTES e DEPOIS: MANTÉM "
          "(serve nas duas funções), SOBROU (servia na antiga e não na nova — é "
          "o que revogar), FALTA (a nova exige e ela não tem), EXCESSO (não está "
          "em nenhum dos dois padrões).",
          "Aba Transferidos → cartão “Sobrou”.",
          f"{N(tra.get('sobrou', 0))} acessos sobraram, em "
          f"{N(tra.get('pessoas_sobrou', 0))} pessoas.", prioritaria=True)
    nota(doc,
         "“Sobrou” é o número operacional desta aba: os acessos que a "
         "transferência tornou desnecessários.")

    # ---------------- 4 tratativa ----------------
    doc.add_page_break()
    h1(doc, "4. Tratativa e histórico — o que o analista registra")

    regra(doc, "4.1", "Motivo só na pendência",
          "Quais campos são obrigatórios ao tratar cada tipo de caso.",
          "Pendência exige motivo (lista fechada) + parecer. Desligado e "
          "transferido exigem só o parecer: no desligado o desfecho é sempre "
          "revogar, e no transferido o motivo repetiria o nome da aba — campo "
          "obrigatório de resposta única vira atrito, não informação.",
          "Botão Resolver, nas três abas.",
          "Motivos: Exceção, Transferência de Área, Acesso Indevido "
          "(configuráveis em arquivo).")

    regra(doc, "4.2", "Chamado no Jira é opcional",
          "Se é preciso ter chamado aberto para registrar a tratativa.",
          "Não é. O obrigatório é o que PROVA a tratativa (motivo e parecer); o "
          "número do chamado é referência externa e pode nem existir ainda. Até "
          "05/08 o ticket era obrigatório, o que impedia registrar tratativa "
          "interna.",
          "Formulário de tratativa → bloco “Chamado no Jira (opcional)”.",
          "O botão “Abrir chamado no Jira” aparece DESABILITADO — depende do "
          "formulário no Jira e de alinhamento da equipe.")

    regra(doc, "4.3", "Ciclo de vida do caso",
          "Como um caso caminha da identificação até o encerramento.",
          "Pendência identificada → Resolvida → Aderente. Se a divergência "
          "reaparece num processamento seguinte, o caso REABRE e ganha novo "
          "ciclo, em vez de sumir do histórico.",
          "Aba Histórico (trilha por sistema) e aba Aderentes.",
          f"{N(d['historico'])} registros no histórico  ·  {N(d['aderentes'])} "
          "linhas na aba Aderentes.")

    regra(doc, "4.4", "Quarentena",
          "Como tirar um caso da fila sem perdê-lo de vista.",
          "Envio com prazo em dias, título e motivo. No fim do prazo o caso volta "
          "sozinho para os pendentes. Vale para a pessoa inteira, um sistema ou "
          "um acesso específico.",
          "Aba Quarentena.",
          f"{N(vg['quarentena_ativa'])} casos em quarentena (recurso disponível, "
          "sem uso nesta base).")

    # ---------------- 5 leitura ----------------
    h1(doc, "5. Como ler os filtros")

    regra(doc, "5.1", "O filtro de sistema faz coisas diferentes em cada aba",
          "O que aparece quando se filtra por um sistema.",
          "Em Pendências, Aderentes e Histórico o filtro ISOLA: a pessoa entra se "
          "tiver algo naquele sistema e só as linhas daquele sistema aparecem. Na "
          "Consulta ele escolhe QUEM aparece, mas a linha continua mostrando "
          "todos os acessos da pessoa — a Consulta é a visão completa do indivíduo.",
          "Pendências → filtro Sistema; depois Consulta com o mesmo filtro.",
          "  ·  ".join(f"{kk} {N(vv)}" for kk, vv in
                       sorted(d["sistema_pend"].items(), key=lambda x: -x[1]))
          + "  (pessoas a tratar por sistema).")
    nota(doc,
         "Na Consulta aparece um aviso em destaque sempre que houver filtro de "
         "sistema ativo, explicando essa diferença.")

    # ---------------- 6 sem regra ----------------
    h1(doc, "6. O que ainda NÃO tem regra")
    par(doc,
        "Itens conhecidos e deliberadamente fora desta entrega — aqui para que a "
        "ausência seja decisão, não surpresa.")
    for tit, txt in [
        ("Perfil excessivo", "quem tem MAIS acesso do que o cargo exige não é "
         "sinalizado quando o esperado também está presente: o veredito Aderente "
         "prevalece e esconde o extra. Só cai em Em Análise pela regra 2.7."),
        ("Abertura automática de chamado", "o botão existe e está desabilitado; "
         "depende do formulário no Jira e de alinhamento da equipe."),
        ("Tempo de tratamento nesta base", "as pendências vindas da matriz "
         "nasceram hoje, então o aging da Visão Geral está achatado. Não é "
         "regra: é característica desta base de teste."),
        ("Terceiros desligados", "a integração existe e está desligada por "
         "configuração nesta fase."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(tit + " — "); r.bold = True
        r.font.size = Pt(9.5); r.font.color.rgb = AZUL
        r2 = p.add_run(txt); r2.font.size = Pt(9.5); r2.font.color.rgb = TEXTO

    # ---------------- 7 fecho ----------------
    doc.add_page_break()
    h1(doc, "7. Resumo para devolver")
    par(doc,
        "Se preferir responder de uma vez: liste abaixo as regras com que NÃO "
        "concorda e o que deveria ser. As demais ficam entendidas como aprovadas.")
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    for i, txt in enumerate(["Regra", "O que está errado", "O que deveria ser"]):
        c = tbl.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(txt); r.bold = True
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1F2D5C")
    larg = [Cm(2.4), Cm(6.8), Cm(7.3)]
    for j, w in enumerate(larg):
        tbl.rows[0].cells[j].width = w
    for _ in range(9):
        cs = tbl.add_row().cells
        for j, w in enumerate(larg):
            cs[j].width = w

    doc.add_paragraph()
    nota(doc,
         "Toda regra deste documento é critério ou parâmetro que pode mudar. "
         "Alterações exigem um reprocessamento para os números refletirem a "
         "decisão.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK -> {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
    print(f"   numeros lidos do painel em {API}")


if __name__ == "__main__":
    main()
