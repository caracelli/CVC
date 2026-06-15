# -*- coding: utf-8 -*-
"""Gera o documento de entrega (.docx) da Fase 1 — SYSTUR.

PUBLICO: pessoa de NEGOCIO (nao tecnica). Linguagem simples, sem jargao,
sem nomes de arquivo .exe/paths/colunas em MAIUSCULA/termos de TI.

Saida: ENTREGA/CVC_IAM_Analytics_Entrega_SYSTUR_v1.0.docx

Uso:  python scripts/gerar_documento_entrega.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

RAIZ = Path(__file__).resolve().parent.parent
SAIDA = RAIZ / "ENTREGA" / "CVC_IAM_Analytics_Entrega_SYSTUR_v1.0.docx"
# Logo OFICIAL da CVC Corp: deixe o arquivo neste caminho que o gerador o insere
# na capa automaticamente. (Sem o arquivo, a capa sai sem logo.)
LOGO = RAIZ / "scripts" / "cvc_logo_oficial.png"
PRINTS = RAIZ / "docs" / "prints_painel"

NAVY = RGBColor(0x1F, 0x2D, 0x5C)
TEAL = RGBColor(0x16, 0xA0, 0x85)
CINZA = RGBColor(0x5A, 0x63, 0x77)
NAVY_HEX = "1F2D5C"
CLARO_HEX = "EAF1F4"


def _shade(cell, hex_cor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_cor)
    tcpr.append(shd)


def _set_cell(cell, texto, *, bold=False, branco=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(str(texto))
    run.font.size = Pt(size)
    run.font.bold = bold
    if branco:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def tabela(doc, cabecalho, linhas, larguras=None):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, c in enumerate(cabecalho):
        cell = t.rows[0].cells[i]
        _set_cell(cell, c, bold=True, branco=True)
        _shade(cell, NAVY_HEX)
    for r, linha in enumerate(linhas):
        cells = t.add_row().cells
        for i, val in enumerate(linha):
            _set_cell(cells[i], val)
            if r % 2 == 1:
                _shade(cells[i], CLARO_HEX)
    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def h1(doc, texto):
    p = doc.add_heading(level=1)
    run = p.add_run(texto)
    run.font.color.rgb = NAVY
    run.font.size = Pt(16)
    return p


def h2(doc, texto):
    p = doc.add_heading(level=2)
    run = p.add_run(texto)
    run.font.color.rgb = TEAL
    run.font.size = Pt(13)
    return p


def par(doc, texto, *, italico=False, cor=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.italic = italico
    if cor:
        run.font.color.rgb = cor
    return p


def code_block(doc, texto):
    """Bloco com fundo claro — para a 'organizacao das pastas'."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6); pf.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), CLARO_HEX)
    pPr.append(shd)
    for i, linha in enumerate(texto.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(linha)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
    return p


def bullet(doc, texto, *, neg=None):
    p = doc.add_paragraph(style="List Bullet")
    if neg:
        r = p.add_run(neg + " ")
        r.font.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(texto)
    r2.font.size = Pt(10.5)
    return p


def imagem(doc, arquivo, width=6.7):
    cam = PRINTS / arquivo
    if not cam.exists():
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(cam), width=Inches(width))


def construir():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---- CAPA ----
    if LOGO.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        p.add_run().add_picture(str(LOGO), width=Inches(2.3))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("CVC IAM Analytics"); r.font.size = Pt(30); r.font.bold = True
    r.font.color.rgb = NAVY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Governança e Análise de Acessos"); r.font.size = Pt(14)
    r.font.color.rgb = CINZA
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Documento de Entrega — Fase 1 (SYSTUR)")
    r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = TEAL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Versão 1.0.0"); r.font.size = Pt(13); r.font.color.rgb = CINZA
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(date.today().strftime("Inclusão e Alteração de Acessos · %d/%m/%Y"))
    r.font.size = Pt(11); r.font.color.rgb = CINZA
    doc.add_page_break()

    # ---- 1. VISAO GERAL ----
    h1(doc, "1. Visão geral")
    par(doc, "O CVC IAM Analytics é a ferramenta que ajuda a CVC a manter os acessos "
             "aos sistemas sob controle. Ela compara o que cada pessoa PODE acessar "
             "(de acordo com o cargo) com o que ela REALMENTE acessa, e mostra, de "
             "forma simples, onde existem diferenças: quem precisa receber um acesso, "
             "quem está com o acesso errado e quem precisa de uma análise.")
    par(doc, "Esta primeira entrega cobre o sistema SYSTUR. Ela é a base sobre a qual "
             "os demais sistemas e as próximas etapas serão acrescentados, sem "
             "necessidade de refazer o que já foi feito.")
    par(doc, "Como ler os números: a ferramenta conta PESSOAS a tratar, não acessos. "
             "Se uma pessoa tem várias pendências, ela aparece como uma pessoa só — o "
             "foco é \"quantas pessoas eu preciso resolver\".", italico=True, cor=NAVY)

    # ---- 2. O QUE A ENTREGA CONTEMPLA ----
    h1(doc, "2. O que esta entrega contempla")
    h2(doc, "2.1. Informações usadas")
    par(doc, "Nesta fase, a ferramenta analisa o sistema SYSTUR. Para isso, ela usa "
             "quatro arquivos, normalmente fornecidos pelas áreas de RH e de TI:")
    tabela(doc,
           ["Arquivo", "Para que serve"],
           [["Funcionários ativos (RH)", "Quem são as pessoas e quais são os seus cargos."],
            ["Mapeamento CCO/CSC", "Liga o centro de custo e o gestor ao acesso esperado."],
            ["Matriz de Perfil de Acesso SYSTUR", "Quais acessos cada cargo deve ter no SYSTUR."],
            ["Extrato do SYSTUR", "Os acessos que as pessoas realmente têm hoje."]],
           larguras=[2.4, 4.3])

    h2(doc, "2.2. O que a ferramenta aponta")
    par(doc, "Cada pessoa recebe uma classificação simples, conforme a comparação "
             "entre o que ela acessa e o que o cargo prevê:")
    tabela(doc,
           ["Classificação", "O que significa", "O que fazer"],
           [["Aderente", "O acesso está correto, conforme o cargo.", "Nada — está tudo certo."],
            ["Incluir Acesso", "O cargo exige um acesso que a pessoa ainda não tem.", "Conceder o acesso."],
            ["Alterar Perfil", "A pessoa tem acesso, mas com o perfil errado.", "Ajustar o perfil."],
            ["Em Análise", "O cargo admite mais de um perfil possível.", "Analisar e decidir."],
            ["Não Mapeado", "Há um acesso de alguém que não está na base de RH ativa.", "Investigar quem é."]],
           larguras=[1.5, 3.5, 1.7])

    h2(doc, "2.3. O que ainda não entra nesta fase")
    par(doc, "Para manter esta primeira entrega objetiva, alguns pontos ficam para as "
             "próximas etapas:")
    bullet(doc, "a retirada de acesso de quem saiu da empresa é tratada em uma etapa separada.", neg="Desligados:")
    bullet(doc, "aguardam a definição da regra de tratamento.", neg="Terceiros:")
    bullet(doc, "os outros sistemas entram nas próximas fases.", neg="Demais sistemas:")

    # ---- 3. OS DOIS PROGRAMAS ----
    h1(doc, "3. Os dois programas")
    par(doc, "A solução tem dois programas que trabalham juntos: um prepara os "
             "resultados e o outro os apresenta.")

    h2(doc, "3.1. O Processador")
    par(doc, "É o motor da análise. Ele faz, de ponta a ponta:")
    bullet(doc, "lê os arquivos colocados na pasta de entrada (RH, matriz e extrato);")
    bullet(doc, "padroniza as informações (acerta maiúsculas, acentos, espaços, códigos) para comparar corretamente;")
    bullet(doc, "cruza o que cada pessoa acessa com o que o cargo prevê e gera a classificação (Aderente, Incluir, Alterar, Em Análise, Não Mapeado);")
    bullet(doc, "grava o resultado na base de dados compartilhada — é essa base que o Visualizador lê;")
    bullet(doc, "arquiva os arquivos já usados e registra um log de cada execução (e separa os que tiveram problema).")
    par(doc, "Quando usar: sob demanda, sempre que chegarem bases novas. Quem usa: a "
             "pessoa responsável pelo processamento. Não precisa ficar aberto — roda, "
             "conclui e fecha.")

    h2(doc, "3.2. O Visualizador (Painel)")
    par(doc, "É o programa do dia a dia. O Visualizador lê a base gerada pelo "
             "Processador e abre o painel no navegador — a parte visual, organizada em "
             "seis telas (detalhadas no item 6). No painel, o time:")
    bullet(doc, "consulta a situação de acesso de qualquer pessoa;")
    bullet(doc, "acompanha os indicadores e a evolução dos atendimentos;")
    bullet(doc, "trata as pendências: resolve (sob um chamado), coloca em quarentena e exporta para Excel.")
    par(doc, "Quem usa: todo o time. Várias pessoas podem abrir ao mesmo tempo, cada "
             "uma no seu computador, vendo o mesmo cenário. Diferente do Processador, "
             "ele fica aberto enquanto a pessoa trabalha.")

    par(doc, "Em resumo: o Processador gera os dados; o Visualizador mostra e trata. "
             "Quem só acompanha precisa apenas do Visualizador.", italico=True, cor=NAVY)

    # ---- 4. VARIAS PESSOAS AO MESMO TEMPO ----
    h1(doc, "4. Como várias pessoas usam ao mesmo tempo")
    par(doc, "A ferramenta foi pensada para uma equipe trabalhar junta, com "
             "tranquilidade:")
    bullet(doc, "os dados ficam em um lugar compartilhado da empresa, acessível a todos do time.")
    bullet(doc, "cada pessoa abre o Visualizador no seu próprio computador e enxerga o mesmo cenário.")
    bullet(doc, "as ações de cada pessoa (tratar uma pendência, colocar um acesso em observação) são guardadas com segurança, sem uma atrapalhar a outra.")
    bullet(doc, "os programas se mantêm atualizados sozinhos (explicado no item 5.4).")

    # ---- 5. COMO UTILIZAR ----
    h1(doc, "5. Como utilizar")

    h2(doc, "5.1. Organização das pastas")
    par(doc, "Ao receber a solução, esta é a organização das pastas. No dia a dia "
             "você usa principalmente duas: a ENTRADA (onde você coloca os arquivos) "
             "e a DADOS (preenchida pela ferramenta com os resultados).")
    code_block(doc,
        "CVC_IAM_ANALYTICS\n"
        "│\n"
        "├── ENTRADA  ← aqui VOCÊ coloca os arquivos\n"
        "│     ├── RH ............. base de funcionários\n"
        "│     ├── MATRIZES ....... Matriz de Perfil de Acesso e Mapeamento CCO/CSC\n"
        "│     └── SISTEMAS ....... extrato de acesso (SYSTUR)\n"
        "│\n"
        "├── DADOS  ← a ferramenta preenche sozinha\n"
        "│     ├── BANCO ......... a base de dados gerada\n"
        "│     ├── SAIDAS ........ relatórios\n"
        "│     ├── PROCESSADOS ... arquivos já usados\n"
        "│     ├── ERROS ......... arquivos com problema\n"
        "│     └── LOGS .......... registros de cada execução\n"
        "│\n"
        "├── EXECUTAVEIS  ← os programas (Processador e Visualizador)\n"
        "│\n"
        "└── INTERACOES  ← ações dos usuários (preenchida automaticamente)")
    par(doc, "Observação: as pastas aparecem com os nomes em maiúsculas exatamente "
             "assim (ENTRADA, DADOS, EXECUTAVEIS, INTERACOES). É só seguir esses nomes "
             "ao guardar cada arquivo.", italico=True, cor=CINZA)

    h2(doc, "5.2. Onde colocar cada arquivo")
    par(doc, "São quatro arquivos. Cada um vai na sua pasta. Podem estar em Excel ou "
             "em CSV. Ao lado, as informações que cada arquivo precisa conter:")
    tabela(doc,
           ["Arquivo", "Onde colocar", "Informações que precisa ter"],
           [["Funcionários ativos (RH)",
             "ENTRADA › RH › ATIVOS",
             "Matrícula, Nome, CPF, Cargo, Departamento, Centro de custo, Data de admissão, E-mail, Situação"],
            ["Mapeamento CCO/CSC",
             "ENTRADA › MATRIZES › ORGANIZACIONAL",
             "Cargo (código e descrição), Departamento, Centro de custo"],
            ["Matriz de Perfil de Acesso SYSTUR",
             "ENTRADA › MATRIZES › PERFIS_SISTEMAS",
             "Cargo, Sistema, Perfil"],
            ["Extrato do SYSTUR",
             "ENTRADA › SISTEMAS › SYSTUR",
             "Usuário, Nome, Perfil, Situação, Data de criação, Último acesso"]],
           larguras=[1.9, 2.1, 3.1])
    par(doc, "Dica: pode colocar mais de um arquivo na mesma pasta — a ferramenta lê "
             "todos. Depois de usados, os arquivos são arquivados automaticamente.",
        italico=True, cor=CINZA)

    h2(doc, "5.3. Passo a passo")
    par(doc, "Importante: a solução já estará instalada na pasta compartilhada da "
             "empresa (instalação única, feita pela TI). Os programas, porém, NÃO "
             "rodam de lá — cada pessoa copia a pasta EXECUTAVEIS para o seu próprio "
             "computador e roda a partir dela. A pasta compartilhada guarda apenas os "
             "dados, que todos enxergam.", italico=True, cor=NAVY)
    passos = [
        ("Copiar os programas para o seu computador (cada pessoa)",
         "Na pasta compartilhada, copie a pasta EXECUTAVEIS para uma pasta no seu "
         "computador (por exemplo, em C:\\CVC) e use a partir dela. Vale tanto para "
         "quem vai processar quanto para quem só vai acompanhar."),
        ("Colocar os arquivos (pessoa responsável)",
         "Guardar os quatro arquivos nas pastas da ENTRADA, na pasta compartilhada, "
         "conforme o item 5.2."),
        ("Processar (pessoa responsável)",
         "Abrir o Processador a partir da sua cópia. Ele compara tudo e grava os "
         "resultados na base compartilhada; arquiva os arquivos usados e separa os que "
         "tiverem problema para conferência."),
        ("Abrir o Visualizador (cada pessoa)",
         "Abrir o Visualizador a partir da sua cópia. Ele abre o painel no navegador, "
         "já com os resultados. A partir daí, é só acompanhar e tratar as pendências "
         "pelas telas (item 6)."),
        ("Atualizar o cenário",
         "Sempre que chegarem bases novas, repetir os passos 2 e 3 (colocar os "
         "arquivos e processar). Mexer apenas no painel não exige processar de novo."),
    ]
    for i, (titulo, desc) in enumerate(passos, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"Passo {i} — {titulo}. ")
        r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
        r2 = p.add_run(desc); r2.font.size = Pt(10.5)

    h2(doc, "5.4. Atualização automática")
    par(doc, "Como cada pessoa usa uma cópia no próprio computador, os programas se "
             "atualizam sozinhos — não é preciso reinstalar em cada máquina quando sai "
             "uma versão nova:")
    bullet(doc, "ao abrir, o programa verifica se existe uma versão mais nova no lugar compartilhado.")
    bullet(doc, "se existir, ele se atualiza sozinho e reabre já na versão nova.")
    bullet(doc, "se o lugar compartilhado estiver indisponível no momento, ele continua funcionando normalmente e tenta atualizar na próxima abertura.")
    par(doc, "Para liberar uma versão nova para todo o time, basta atualizar os "
             "programas no lugar compartilhado — cada computador se atualiza sozinho na "
             "próxima vez que abrir.", italico=True, cor=CINZA)

    # ---- 6. AS TELAS DO PAINEL ----
    h1(doc, "6. As telas do Painel")
    par(doc, "O Painel tem seis telas. Abaixo, o que cada uma mostra, o que dá para "
             "fazer e como usar. As imagens são ilustrativas, com dados reais do "
             "ambiente.")
    par(doc, "Em todas as listas você pode: buscar por texto, filtrar por qualquer "
             "coluna (clicando no título da coluna), ordenar, escolher quantos itens "
             "ver por página e exportar para Excel — e a exportação sai igual ao que "
             "está na tela, respeitando os filtros aplicados.", italico=True, cor=CINZA)
    telas = [
        ("Visão Geral", "01_visao_geral.png",
         "É o resumo gerencial do dia a dia. Não tem ações — é uma tela de leitura "
         "para enxergar o todo e priorizar.",
         "Mostra:",
         ["quantas pessoas há em cada classificação (Incluir, Alterar, Em Análise, "
          "Não Mapeado, Aderente);",
          "a evolução dos atendimentos nos últimos 30 dias: identificados, resolvidos "
          "e regularizados;",
          "o tempo de tratamento do ciclo (Pendência → Resolvido → Aderente);",
          "a distribuição das pendências por tipo e por sistema;",
          "o tempo que as pendências estão abertas (aging) e a movimentação de RH."]),
        ("Consulta", "02_consulta.png",
         "Para investigar uma pessoa específica e ver a situação completa dela.",
         "Você pode:",
         ["buscar por nome, matrícula, login, CPF ou e-mail;",
          "expandir a pessoa para ver todos os acessos e pendências dela;",
          "abrir um painel lateral com os detalhes;",
          "saltar direto para as telas de Pendências, Histórico ou Aderentes daquela "
          "pessoa, pelos atalhos."]),
        ("Pendências", "03_pendencias.png",
         "É a lista de trabalho — por pessoa, o que precisa ser incluído, alterado ou "
         "analisado. É aqui que o tratamento acontece.",
         "Você pode:",
         ["filtrar por qualquer coluna (vínculo, departamento, status, tipo…) e por "
          "período;",
          "agrupar por pessoa e expandir/recolher os acessos (botão +/−);",
          "resolver uma pendência, informando o chamado e uma descrição;",
          "enviar um acesso para quarentena (observação por prazo);",
          "salvar filtros favoritos para reusar; exportar para Excel igual à tela."]),
        ("Aderentes", "04_aderentes.png",
         "As pessoas que já estão com os acessos corretos (conformes). Use para "
         "conferir o que já foi resolvido.",
         "Mostra / você pode:",
         ["a trilha de datas do ciclo (quando virou pendência, quando foi resolvido, "
          "quando ficou aderente) e o tempo total até regularizar;",
          "buscar e filtrar por coluna (cargo, sistema, perfil…);",
          "exportar a lista para Excel."]),
        ("Histórico", "05_historico.png",
         "A trilha do que aconteceu ao longo do tempo. Serve como evidência para "
         "auditoria.",
         "Mostra / você pode:",
         ["admissões, mudanças de cargo/área e a evolução de cada acesso "
          "(pendência → resolvido → aderente);",
          "filtrar por período e buscar por matrícula ou nome;",
          "expandir cada pessoa para ver a sequência completa; exportar para Excel."]),
        ("Quarentena", "06_quarentena.png",
         "Acessos colocados em observação por um prazo, antes de uma decisão "
         "definitiva. Use quando precisar de um tempo para avaliar antes de agir.",
         "Você pode:",
         ["ver as quarentenas ativas, com os dias restantes de cada uma;",
          "consultar o histórico de quarentenas já encerradas;",
          "filtrar e exportar para Excel."]),
    ]
    for nome, arq, desc, lead, acoes in telas:
        doc.add_page_break()          # cada tela na sua pagina (titulo junto do print)
        h2(doc, nome)
        par(doc, desc)
        if lead:
            p = doc.add_paragraph()
            r = p.add_run(lead); r.font.bold = True; r.font.size = Pt(10.5)
            r.font.color.rgb = NAVY
        for a in acoes:
            bullet(doc, a)
        imagem(doc, arq)

    # ---- 7. ATENDIMENTO ----
    h1(doc, "7. Como funciona o atendimento")
    par(doc, "Cada pendência percorre três etapas, e a ferramenta registra a data de "
             "cada uma (a trilha completa fica na tela de Histórico):")
    tabela(doc,
           ["Etapa", "Quando acontece"],
           [["1. Identificada", "A ferramenta encontra a diferença de acesso."],
            ["2. Resolvida", "A pessoa responsável trata o caso (sob um chamado) e marca como resolvido."],
            ["3. Regularizada", "Em uma próxima análise, o acesso passa a estar correto."]],
           larguras=[1.7, 5.0])
    par(doc, "Pontos importantes:")
    bullet(doc, "o tratamento é feito por pessoa, sempre vinculado a um chamado, com uma breve descrição do que foi feito.")
    bullet(doc, "a ferramenta mede quanto tempo leva cada etapa, ajudando a acompanhar o ritmo de atendimento.")
    bullet(doc, "tudo o que é exportado para Excel sai igual ao que está na tela, respeitando os filtros aplicados.")

    # ---- 8. PROXIMAS ETAPAS ----
    h1(doc, "8. O que ainda será entregue")
    par(doc, "As próximas etapas se somam a esta entrega. A sequência prevista é:")
    tabela(doc,
           ["Etapa", "Entrega", "Quando"],
           [["Demais sistemas", "Inclusão dos outros sistemas, um a um", "junho/2026"],
            ["Visão unificada", "Os acessos de todos os sistemas em uma visão só", "junho/2026"],
            ["Regras completas", "Matrizes e regras de acesso de todos os sistemas", "julho/2026"],
            ["Desligados", "Retirada de acesso de quem saiu da empresa", "julho/2026"],
            ["Transferidos", "Revisão de acesso após mudança de área", "jul-ago/2026"],
            ["Chamados automáticos", "Abertura automática de chamados no Jira", "agosto/2026"]],
           larguras=[1.6, 3.7, 1.4])
    par(doc, "Quando um novo sistema entra, ele é simplesmente \"ligado\" na "
             "ferramenta — o Painel já está preparado para vários sistemas.",
        italico=True, cor=CINZA)

    # ---- 9. SUGESTOES ----
    h1(doc, "9. Sugestões e recomendações")
    par(doc, "Alguns pontos que recomendamos combinar para aproveitar melhor a "
             "ferramenta e preparar as próximas fases:")
    bullet(doc, "definir de quanto em quanto tempo as bases serão atualizadas e analisadas (por exemplo, uma vez por mês), para ter um ciclo previsível.", neg="Periodicidade:")
    bullet(doc, "definir quem é responsável por tratar cada tipo de pendência e em quanto tempo (prazo-alvo).", neg="Responsáveis e prazos:")
    bullet(doc, "padronizar o uso do chamado em todo tratamento — isso já prepara a abertura automática de chamados (etapa futura).", neg="Uso do chamado:")
    bullet(doc, "combinar o prazo e o critério para um acesso sair da observação, e quem decide.", neg="Quarentena:")
    bullet(doc, "manter uma rotina de cópia de segurança da base, que guarda todo o histórico.", neg="Cópia de segurança:")
    bullet(doc, "o histórico serve como evidência para auditorias (inclusive LGPD); vale definir por quanto tempo guardar.", neg="Evidência para auditoria:")
    bullet(doc, "revisar periodicamente as matrizes de perfil por cargo, pois elas definem o que é \"esperado\" — uma matriz desatualizada gera diferenças que não existem de fato.", neg="Revisão das matrizes:")
    bullet(doc, "a tela de Visão Geral vai evoluir conforme o uso e o feedback do time — sugestões são bem-vindas.", neg="Evolução do Painel:")

    # ---- 10. GLOSSARIO ----
    h1(doc, "10. Glossário")
    tabela(doc,
           ["Termo", "O que quer dizer"],
           [["Aderente", "Pessoa cujo acesso está correto, conforme o cargo."],
            ["Pendência", "Uma diferença de acesso que precisa ser tratada."],
            ["Em Análise", "Cargo que admite mais de um perfil — precisa de decisão."],
            ["Não Mapeado", "Acesso de alguém que não está na base de RH ativa."],
            ["Quarentena", "Acesso colocado em observação por um prazo antes da decisão."],
            ["Perfil", "O conjunto de permissões que a pessoa tem dentro de um sistema."],
            ["Matriz de perfil", "Tabela que define qual perfil cada cargo deve ter em cada sistema."],
            ["Centro de custo", "Código que identifica a área/unidade à qual a pessoa pertence."]],
           larguras=[1.8, 5.3])

    par(doc, "")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CVC IAM Analytics · Versão 1.0.0 · Documento de entrega — Fase 1 (SYSTUR)")
    r.font.size = Pt(8); r.font.color.rgb = CINZA

    SAIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SAIDA)
    print(f"OK -> {SAIDA}")
    print(f"  {SAIDA.stat().st_size/1024:.0f} KB")


if __name__ == "__main__":
    construir()
