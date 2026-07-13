# -*- coding: utf-8 -*-
"""Gera docs/DOCUMENTACAO_PAINEL.docx — versao curta.

Foco: o que cada filtro faz + regras de cada aba (Pendencias, Quarentena,
Historico). Sem detalhes tecnicos, sem glossario.

Uso:
    pip install python-docx
    python scripts/gerar_doc_painel.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(__file__).resolve().parent.parent / "docs" / "DOCUMENTACAO_PAINEL.docx"

COR_AZUL = RGBColor(0x1F, 0x2D, 0x5C)
COR_CINZA = RGBColor(0x5A, 0x64, 0x78)
COR_TEXTO = RGBColor(0x2C, 0x33, 0x40)


def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for r in p.runs:
        r.font.color.rgb = COR_AZUL


def p(doc, texto, italic=False, bold=False, size=10):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = COR_TEXTO
    return par


def tabela(doc, headers, rows, larguras_cm=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        par = hdr[i].paragraphs[0]
        r = par.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade(hdr[i], "1F2D5C")
        if larguras_cm and i < len(larguras_cm):
            hdr[i].width = Cm(larguras_cm[i])
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            par = cells[ci].paragraphs[0]
            r = par.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = COR_TEXTO
            if ri % 2 == 1:
                shade(cells[ci], "F5F6FA")
            if larguras_cm and ci < len(larguras_cm):
                cells[ci].width = Cm(larguras_cm[ci])
    doc.add_paragraph()
    return tbl


def main():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(2)
        sec.top_margin = sec.bottom_margin = Cm(2)

    # Capa
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("CVC · IAM ANALYTICS")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = COR_AZUL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Guia rápido do painel")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = COR_AZUL
    doc.add_paragraph()

    # ── Filtros ────────────────────────────────────────────────────────────
    h1(doc, "Filtros laterais")
    p(doc,
      "Aparecem em todas as abas. Clique simples isola o valor; "
      "Ctrl+clique adiciona/remove (multi-seleção).")
    tabela(doc,
        ["Filtro", "Para que serve"],
        [
            ["Vínculo",
             "Separa Funcionário (CLT, na base RH) de Terceiro "
             "(quando a integração existir)."],
            ["Ação",
             "O que precisa ser feito com a pendência: Incluir Acesso, "
             "Alterar Perfil, Em Análise ou Usuário Não Encontrado."],
            ["Status",
             "Pendente (ainda não tratado) ou Resolvido (já passou pela "
             "Resolução)."],
            ["Tipo",
             "Mesma classificação da Ação, com o rótulo usado na grid: "
             "Sem Acesso, Divergente, Em Análise, Sem Vínculo RH."],
            ["Sistema",
             "Filtra por sistema de origem do acesso. Na Fase 1, apenas "
             "SYSTUR."],
        ],
        larguras_cm=[3, 13.5])

    # ── O que cada Ação significa ─────────────────────────────────────────
    h1(doc, "O que cada Ação significa")
    tabela(doc,
        ["Ação", "Quando aparece"],
        [
            ["Incluir Acesso",
             "Funcionário sem o acesso que a matriz do cargo exige."],
            ["Alterar Perfil",
             "Funcionário com perfil diferente do permitido para o cargo."],
            ["Em Análise",
             "O cargo tem mais de um perfil possível na matriz — alguém "
             "precisa decidir."],
            ["Usuário Não Encontrado",
             "Acesso no sistema sem funcionário correspondente no RH "
             "ativo (usuário \"órfão\")."],
        ],
        larguras_cm=[3.5, 13])

    # ── Aba Pendências ────────────────────────────────────────────────────
    h1(doc, "Aba Pendências")
    p(doc, "Lista todos os funcionários (e usuários sem vínculo) com "
           "alguma pendência de acesso. Uma linha por pessoa.")
    p(doc, "Regras aplicadas:", bold=True)
    p(doc, "• Mostra somente pendências com Ação (Incluir, Alterar, "
           "Em Análise ou Usuário Não Encontrado).")
    p(doc, "• A coluna Status indica se já houve resolução: "
           "Pendente = não, Resolvido = sim.")
    p(doc, "• A linha agrupa todas as pendências do mesmo funcionário "
           "(quando há mais de uma, o número aparece na coluna Qtd e dá "
           "pra expandir).")
    p(doc, "Ações disponíveis na linha:", bold=True)
    tabela(doc,
        ["Botão", "O que faz"],
        [
            ["Resolver (⊕)",
             "Abre um modal pra registrar a resolução sob ticket do Jira. "
             "Depois de confirmar, a linha passa a Resolvido e aparece no "
             "Histórico."],
            ["Quarentena",
             "Coloca o funcionário em quarentena por 90 dias. Ele sai "
             "das Pendências e vai pra aba Quarentena."],
            ["Lupa (🔍)",
             "Aparece quando a pendência já foi resolvida. Abre o modal "
             "com os detalhes da resolução (ticket, descrição, quem "
             "resolveu)."],
        ],
        larguras_cm=[3.5, 13])

    # ── Aba Quarentena ────────────────────────────────────────────────────
    h1(doc, "Aba Quarentena")
    p(doc, "Funcionários colocados em \"compasso de espera\" antes da "
           "decisão final.")
    p(doc, "Regras aplicadas:", bold=True)
    p(doc, "• A quarentena dura 90 dias a partir do envio (encerra "
           "sozinha após esse prazo).")
    p(doc, "• Quem manda pra cá é o usuário do painel (botão Quarentena "
           "na grid de Pendências).")
    p(doc, "• Dois sub-modos no topo da aba:")
    p(doc, "    – Ativas: quem está em quarentena agora.")
    p(doc, "    – Histórico: quem já saiu (com motivo e data de saída).")
    p(doc, "• Cada linha mostra quem criou a quarentena (usuário Windows) "
           "e quem encerrou (no Histórico).")
    p(doc, "• Botão Retirar (nas Ativas) encerra a quarentena antes do "
           "prazo. Vai pro Histórico com motivo \"Resolvido\".")

    # ── Aba Histórico ─────────────────────────────────────────────────────
    h1(doc, "Aba Histórico")
    p(doc, "Trilha de auditoria das resoluções de pendências. Mostra "
           "o ciclo de vida de cada pendência que foi tratada.")
    p(doc, "Regras aplicadas:", bold=True)
    p(doc, "• Cada resolução gera DUAS linhas:")
    p(doc, "    – Pendência identificada: quando a pendência foi detectada "
           "pelo Processador.")
    p(doc, "    – Pendência resolvida: quando o usuário registrou a "
           "resolução sob ticket do Jira.")
    p(doc, "• A Pendência resolvida é SEMPRE posterior à Pendência "
           "identificada (regra cronológica).")
    p(doc, "• Botão Exportar Excel: gera planilha com o mesmo "
           "agrupamento e formatação da grid.")
    p(doc, "• Histórico apenas lê — nada é editado aqui. As resoluções "
           "vêm da aba Pendências.")

    # Rodape
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("CVC IAM Analytics — v2.0.1")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = COR_CINZA

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}  ({OUT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
